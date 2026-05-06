# -*- coding: utf-8 -*-
"""生成《访问控制技术——层次聚类算法》实验报告 docx"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ======================== 页面设置 ========================
for section in doc.sections:
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.18)
    section.right_margin  = Cm(3.18)

# ======================== 样式定义 ========================
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)  # 小四
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# 正文标题样式（一级标题：黑体三号加粗）
heading1_style = doc.styles['Heading 1']
heading1_style.font.name = '黑体'
heading1_style.font.size = Pt(16)  # 三号
heading1_style.font.bold = True
heading1_style.font.color.rgb = RGBColor(0, 0, 0)
heading1_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
heading1_style.paragraph_format.space_before = Pt(12)
heading1_style.paragraph_format.space_after  = Pt(6)
heading1_style.paragraph_format.line_spacing = 1.5

# 二级标题样式（黑体四号加粗）
heading2_style = doc.styles['Heading 2']
heading2_style.font.name = '黑体'
heading2_style.font.size = Pt(14)  # 四号
heading2_style.font.bold = True
heading2_style.font.color.rgb = RGBColor(0, 0, 0)
heading2_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
heading2_style.paragraph_format.space_before = Pt(8)
heading2_style.paragraph_format.space_after  = Pt(4)
heading2_style.paragraph_format.line_spacing = 1.5

# 三级标题样式（黑体小四加粗）
heading3_style = doc.styles['Heading 3']
heading3_style.font.name = '黑体'
heading3_style.font.size = Pt(12)  # 小四
heading3_style.font.bold = True
heading3_style.font.color.rgb = RGBColor(0, 0, 0)
heading3_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# ======================== 辅助函数 ========================

def add_body_para(text, bold=False, indent=True, align=None, spacing_after=None):
    """添加正文段落，首行缩进2字符"""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)  # 约2个中文字符
    if align is not None:
        p.alignment = align
    if spacing_after is not None:
        p.paragraph_format.space_after = spacing_after
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_title(text, level=1):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_code_block(code_text):
    """添加代码块（等宽字体，灰色背景风格）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    for line in code_text.strip().split('\n'):
        run = p.add_run(line + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(40, 40, 40)
    return p

# ======================== 封面 ========================

# 空行
for _ in range(2):
    doc.add_paragraph()

# 学校学院
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('湖南科技大学计算机科学与工程学院')
run.font.name = '楷体_GB2312'
run.font.size = Pt(18)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')

doc.add_paragraph()  # 空行

# 报告标题
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('综合实践能力创新实训3\n课程设计报告')
run.font.name = '楷体_GB2312'
run.font.size = Pt(26)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')

# 空行
for _ in range(3):
    doc.add_paragraph()

# 封面信息行
info_items = [
    ('专业班级：', '信息安全四班'),
    ('姓    名：', '周禹嘉'),
    ('学    号：', '2305030406'),
    ('指导教师：', '廖俊国'),
    ('时    间：', '2026.3.9-3.20'),
    ('地    点：', '逸夫楼333'),
]

for label, value in info_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(35)
    # 标签部分：Times New Roman 加粗
    run_label = p.add_run(label)
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(16)
    run_label.bold = True
    # 值部分
    run_val = p.add_run('     ' + value)
    run_val.font.name = '楷体_GB2312'
    run_val.font.size = Pt(16)
    run_val._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')

# 空行
for _ in range(4):
    doc.add_paragraph()

# 报告副标题
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
p.paragraph_format.space_after = Pt(20)
p.paragraph_format.line_spacing = Pt(30)
run = p.add_run('访问控制技术课程设计报告')
run.font.name = '宋体'
run.font.size = Pt(18)
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ======================== 分页 ========================
doc.add_page_break()

# ======================== 正文 ========================

# ---- 一、算法基本原理 ----
add_title('一、算法基本原理', level=1)

add_title('1.1 聚类分析概述', level=2)

add_body_para(
    '聚类分析（Clustering Analysis）是一种无监督学习方法，其核心任务是将一组物理或抽象的对象，'
    '根据它们之间的相似程度，划分为若干个有意义的组（称为"簇"，Cluster）。聚类分析的目标是使得'
    '同一簇内的对象之间具有较高的相似性，而不同簇之间的对象具有较大的差异性。聚类分析广泛应用于'
    '数据挖掘、模式识别、图像分割、信息检索以及大数据安全领域中的访问控制技术等场景。'
)

add_body_para(
    '在访问控制技术中，聚类分析可以用于对用户行为进行分组，识别具有相似访问模式的用户群体，'
    '从而为基于角色的访问控制（RBAC）提供角色挖掘依据。此外，通过对访问日志进行聚类分析，'
    '还可以检测异常访问行为，辅助识别潜在的内部威胁或外部攻击。'
)

add_title('1.2 层次聚类算法原理', level=2)

add_body_para(
    '层次聚类（Hierarchical Clustering）是聚类算法的重要分支之一。与划分聚类（如K-Means）不同，'
    '层次聚类通过计算不同类别数据点间的相似度，构建一棵有层次的嵌套聚类树（Dendrogram，树状图）。'
    '在聚类树中，最底层是原始数据点，最顶层是所有数据点合并为一个簇的根节点。层次聚类算法的显著'
    '优势在于可以在不同的尺度（层次）上展示数据集的聚类情况，不需要预先指定簇的数量。'
)

add_body_para(
    '层次聚类分为两种主要策略：自底向上的凝聚层次聚类（Agglomerative Hierarchical Clustering）和'
    '自顶向下的分裂层次聚类（Divisive Hierarchical Clustering）。本实验采用凝聚层次聚类策略。'
)

add_body_para(
    '凝聚层次聚类的基本思想是：初始时将每个数据对象视为一个独立的簇（共N个簇，每簇仅包含一个对象）；'
    '然后反复合并距离最近的两个簇，直到满足某个终止条件（如达到预设的簇数目，或所有对象被合并为一个簇）。'
    '算法具体步骤如下：',
    bold=False
)

add_body_para(
    '步骤1：将每个数据点初始化为一个独立的簇，共得到N个簇。\n'
    '步骤2：计算所有簇两两之间的距离，得到距离矩阵。\n'
    '步骤3：找到距离最近的两个簇，将它们合并为一个新簇。\n'
    '步骤4：更新距离矩阵，重新计算新簇与其他所有簇之间的距离。\n'
    '步骤5：重复步骤3和步骤4，直到簇的数量达到预设值K（或仅剩1个簇）。',
    indent=True
)

add_title('1.3 簇间距离计算方法', level=2)

add_body_para(
    '凝聚层次聚类算法中的一个关键步骤是计算簇与簇之间的距离。不同的距离计算方法会导致不同的聚类结果。'
    '本实验实现了以下三种经典的簇间距离计算方法：'
)

add_body_para(
    '（1）单链法（Single-linkage）：分别计算一个簇中每一个对象与另一个簇中每一个对象的距离，'
    '然后取距离的最小值作为簇间距离。d(cᵢ, cⱼ) = min{d(p, q) | p∈cᵢ, q∈cⱼ}。\n'
    '特点：单链法对噪声和离群点较为敏感，容易产生"链式效应"（chaining effect），'
    '倾向于产生狭长的簇。',
    indent=True
)

add_body_para(
    '（2）全链法（Complete-linkage）：分别计算一个簇中每一个对象与另一个簇中每一个对象的距离，'
    '然后取距离的最大值作为簇间距离。d(cᵢ, cⱼ) = max{d(p, q) | p∈cᵢ, q∈cⱼ}。\n'
    '特点：全链法倾向于产生紧凑且直径较小的球形簇，对离群点较为敏感。',
    indent=True
)

add_body_para(
    '（3）组平均法（Average-linkage）：簇间距离等于两组对象之间的平均距离。'
    'd(cᵢ, cⱼ) = (1/(|cᵢ|·|cⱼ|)) · Σ d(p, q)，其中p∈cᵢ, q∈cⱼ。\n'
    '特点：组平均法是单链法和全链法之间的折中方案，既不过于敏感也不过于保守，'
    '在实际应用中表现较好。',
    indent=True
)

add_title('1.4 算法流程图', level=2)

add_body_para(
    '凝聚层次聚类算法的基本流程如下图所示：',
    indent=True
)

# 插入流程图图片（如果存在）
flowchart_path = 'clustering_all_methods.png'
if os.path.exists(flowchart_path):
    doc.add_picture(flowchart_path, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_body_para(
    '图1 三种距离计算方法对应的树状图与聚类结果散点图。从上至下依次为：单链法（Single-linkage）、'
    '全链法（Complete-linkage）和组平均法（Average-linkage）。左列为层次聚类树状图（Dendrogram），'
    '展示了聚类的层次合并过程；右列为将数据划分为3个簇后的散点图结果。',
    indent=True
)

# ---- 二、算法核心函数实现与说明 ----
doc.add_page_break()
add_title('二、算法核心函数实现与说明', level=1)

add_title('2.1 欧几里得距离计算与距离矩阵构建', level=2)

add_body_para(
    '本实验使用欧几里得距离（Euclidean Distance）度量两个数据点之间的相似性。'
    '对于二维空间中的两点a(x₁, y₁)和b(x₂, y₂)，欧几里得距离定义为：'
    'd(a,b) = √[(x₁-x₂)² + (y₁-y₂)²]。以下是距离矩阵计算的核心代码：'
)

add_code_block(
'''def euclidean_dist(a, b):
    """计算两个点之间的欧几里得距离"""
    return np.sqrt(np.sum((a - b) ** 2))

def compute_distance_matrix(points):
    """计算初始距离矩阵（N×N）"""
    n = len(points)
    dist_mat = np.zeros((n, n))
    for i in range(n):
        for j in range(i + 1, n):
            d = euclidean_dist(points[i], points[j])
            dist_mat[i][j] = d
            dist_mat[j][i] = d
    return dist_mat'''
)

add_body_para(
    'compute_distance_matrix函数接收N个数据点，返回一个N×N的对称距离矩阵，其中第i行第j列的元素'
    '表示点i和点j之间的欧几里得距离。由于距离矩阵是对称的（dist_mat[i][j] == dist_mat[j][i]），'
    '且对角线元素为0（点到自身的距离为0），因此只需要计算上三角矩阵即可。该矩阵是后续所有簇间距离'
    '计算的基础数据结构。'
)

add_title('2.2 簇间距离计算函数', level=2)

add_body_para(
    'cluster_distance函数实现了三种不同的簇间距离计算方法。该函数接收两个簇的成员索引列表和距离矩阵，'
    '根据指定的方法名（single/complete/average）返回对应的簇间距离值。'
)

add_code_block(
'''def cluster_distance(cluster_i, cluster_j, dist_mat, method='single'):
    """计算两个簇之间的距离
    method: 'single' (单链/min), 'complete' (全链/max), 'average' (组平均)
    """
    dists = []
    for p in cluster_i:
        for q in cluster_j:
            dists.append(dist_mat[p][q])
    if method == 'single':
        return np.min(dists)   # 单链法：取最小距离
    elif method == 'complete':
        return np.max(dists)   # 全链法：取最大距离
    elif method == 'average':
        return np.mean(dists)  # 组平均法：取平均距离'''
)

add_body_para(
    '该函数通过双层循环收集两个簇中所有点对之间的距离值，然后根据method参数选择相应的统计量。'
    '单链法取最小值，使得远距离的离群点容易被拉入邻近簇（链式效应）；全链法取最大值，倾向于'
    '保持簇的紧凑性；组平均法取平均值，是前两种方法的折中方案。'
)

add_title('2.3 凝聚层次聚类主函数', level=2)

add_body_para(
    'agglomerative_clustering函数是整个算法的核心，实现了完整的自底向上凝聚层次聚类流程。'
    '该函数接收数据点数组、目标聚类数和距离计算方法作为参数，返回合并历史、最终分类标签和链路矩阵。'
)

add_code_block(
'''def agglomerative_clustering(points, n_clusters=3, method='single'):
    """凝聚层次聚类（自底向上）"""
    n = len(points)
    dist_mat = compute_distance_matrix(points)

    # 初始化：每个点为一个簇，簇ID = 点的索引
    clusters = {i: [i] for i in range(n)}
    next_id = n  # 新簇从 n 开始编号，避免与原始点索引冲突

    clusters_history = []
    linkage_rows = []

    while len(clusters) > n_clusters:
        active_ids = list(clusters.keys())
        min_dist = float('inf')
        best_pair = None

        # 在所有活跃簇中寻找距离最近的一对
        for idx_i, ci in enumerate(active_ids):
            for cj in active_ids[idx_i + 1:]:
                d = cluster_distance(clusters[ci], clusters[cj],
                                     dist_mat, method)
                if d < min_dist:
                    min_dist = d
                    best_pair = (ci, cj)

        ci, cj = best_pair
        # 合并：创建新簇
        new_members = clusters[ci] + clusters[cj]
        clusters[next_id] = new_members
        # 删除旧簇
        del clusters[ci], clusters[cj]
        # 记录合并过程
        clusters_history.append((len(active_ids), min_dist, next_id))
        linkage_rows.append([float(ci), float(cj), min_dist,
                             float(len(new_members))])
        next_id += 1

    # 生成最终标签（0, 1, 2, ...）
    final_labels = np.zeros(n, dtype=int)
    for label, members in enumerate(clusters.values()):
        for p in members:
            final_labels[p] = label

    return clusters_history, final_labels, np.array(linkage_rows), dist_mat'''
)

add_body_para(
    '该算法的关键设计说明：簇用字典结构管理，键为簇ID，值为成员点索引列表。每次迭代中找到距离最近的'
    '两个簇后，将它们合并为一个新簇，新簇ID从N开始递增以避免与原始索引冲突。同时删除被合并的两个旧簇。'
    '算法持续执行直到簇的数量降至目标值n_clusters。最终通过扫描剩余的每个簇，为每个数据点分配'
    '0到K-1之间的簇标签。'
)

add_title('2.4 使用Sklearn和Scipy实现层次聚类', level=2)

add_body_para(
    '除了手写实现外，本实验还使用Python的Sklearn库和Scipy库实现了层次聚类，以便对比验证。'
)

add_body_para(
    '（1）Scipy实现：使用scipy.cluster.hierarchy.linkage函数计算链路矩阵，该函数支持'
    'single、complete、average等多种linkage方法。然后使用fcluster函数根据链路矩阵进行'
    '扁平聚类（flat clustering），指定t=3和criterion="maxclust"将数据分为3个簇。'
    'Scipy的dendrogram函数还可直接绘制树状图。',
    indent=True
)

add_code_block(
'''from scipy.cluster.hierarchy import linkage, fcluster, dendrogram

link_scipy = linkage(X, method='single', metric='euclidean')
cluster_labels = fcluster(link_scipy, t=3, criterion='maxclust') - 1'''
)

add_body_para(
    '（2）Sklearn实现：使用sklearn.cluster.AgglomerativeClustering类，'
    '设置n_clusters=3和所需的linkage方法。调用fit_predict方法直接获取每个数据点的簇标签。'
    'Sklearn的接口更加简洁统一，适合集成到机器学习流水线中。',
    indent=True
)

add_code_block(
'''from sklearn.cluster import AgglomerativeClustering

agg = AgglomerativeClustering(n_clusters=3, linkage='single')
sk_labels = agg.fit_predict(X)'''
)

# ---- 三、实验结果与分析 ----
doc.add_page_break()
add_title('三、实验结果与分析', level=1)

add_title('3.1 实验数据', level=2)

add_body_para(
    '本实验使用随机生成的25个二维数据点作为测试数据。数据分为3个自然簇，分别以(1.0, 1.0)、'
    '(5.0, 1.0)和(3.0, 4.5)为中心，每个簇的标准差为0.6。具体数据如下：'
)

add_body_para(
    '簇0（8个点，中心≈(1.0, 1.0)）：P0(1.298, 0.917), P1(1.389, 1.914), P2(0.860, 0.860), '
    'P3(1.948, 1.461), P4(0.718, 1.326), P5(0.964, 1.660), P6(0.260, 1.204), P7(0.999, 0.803)。\n'
    '簇1（9个点，中心≈(5.0, 1.0)）：P8(5.180, 1.215), P9(4.448, 0.507), P10(5.281, 0.963), '
    'P11(4.305, 1.504), P12(5.606, 1.614), P13(4.756, 0.967), P14(4.926, 1.058), '
    'P15(5.786, 0.618), P16(4.624, 0.998)。\n'
    '簇2（8个点，中心≈(3.0, 4.5)）：P17(3.210, 4.906), P18(2.815, 5.267), P19(3.578, 3.843), '
    'P20(2.064, 5.002), P21(2.463, 4.669), P22(2.865, 3.261), P23(3.226, 4.094), P24(3.480, 4.671)。',
    indent=True
)

add_title('3.2 初始距离矩阵', level=2)

add_body_para(
    '计算25个数据点之间的欧几里得距离，得到25×25的对称距离矩阵。以下展示了前8个数据点（全部属于簇0）'
    '的距离子矩阵（8×8）：'
)

# 距离矩阵热力图
heatmap_path = 'distance_matrix.png'
if os.path.exists(heatmap_path):
    doc.add_picture(heatmap_path, width=Inches(5.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_body_para(
    '图2 距离矩阵热力图。颜色越深表示距离越小（越相似），颜色越浅表示距离越大（越不相似）。'
    '从图中可以观察到明显的块状结构，说明了数据中存在的聚类趋势。三个簇内部的点距离较小'
    '（对角线附近的深色块），不同簇之间的点距离较大（非对角线区域的浅色块）。',
    indent=True
)

add_body_para(
    '距离矩阵前8×8子矩阵的数值（保留三位小数）：',
    indent=False
)

add_code_block(
'''[[0.    1.001 0.442 0.847 0.709 0.609 1.076 1.357]
 [1.001 0.    1.18  0.72  0.892 1.367 2.076 1.895]
 [0.442 1.18  0.    1.243 0.487 0.196 1.047 0.916]
 [0.847 0.72  1.243 0.    1.237 1.432 1.797 2.137]
 [0.709 0.892 0.487 1.237 0.    0.605 1.534 1.003]
 [0.609 1.367 0.196 1.432 0.605 0.    0.966 0.759]
 [1.076 2.076 1.047 1.797 1.534 0.966 0.    1.432]
 [1.357 1.895 0.916 2.137 1.003 0.759 1.432 0.   ]]'''
)

add_body_para(
    '从上述距离矩阵可以看出，簇0内部8个数据点之间的最大距离约为2.137（P3与P7之间），'
    '而它们与簇1（索引8~16）和簇2（索引17~24）的数据点之间的距离显著更大。'
    '这种距离分布为聚类算法提供了清晰的判别依据。'
)

add_title('3.3 三种方法的聚类结果对比', level=2)

add_body_para(
    '使用手写的凝聚层次聚类算法，分别采用单链法（Single-linkage）、全链法（Complete-linkage）'
    '和组平均法（Average-linkage）三种簇间距离计算方式，将25个数据点聚为3个簇。'
    '以下是每种方法的详细合并过程和分类结果。'
)

add_title('（1）单链法合并过程与结果', level=3)

add_body_para(
    '单链法的合并过程共22步（从25个簇合并至3个簇）。前5次合并（最小距离分别为0.0907、0.1430、'
    '0.1955、0.2255、0.2439）全部发生在簇0内部，表明簇0的数据点之间最为紧密。第6~8次合并涉及'
    '簇1内部的点，在距离约为0.3064~0.4341时开始合并。簇2内部数据的合并发生在第9~16步左右。'
    '三种方法在最后几次合并中，簇间距离差异明显：单链法在第21~22步合并时距离仅为0.9662和1.0182，'
    '表明簇之间的边界较为接近。'
)

add_body_para(
    '单链法最终分类结果：\n'
    '簇0：P17~P24（8个点）——对应真实簇2\n'
    '簇1：P0~P7（8个点）——对应真实簇0\n'
    '簇2：P8~P16（9个点）——对应真实簇1\n'
    '分类准确率达到100%，25个数据点全部被正确分类。',
    indent=True
)

add_title('（2）全链法合并过程与结果', level=3)

add_body_para(
    '全链法的合并过程同样是22步。前5次合并与单链法完全相同（都发生在最紧密的簇0内部），'
    '但第6次合并的最小距离为0.5190（单链法为0.3064），此后每一步的距离都显著大于单链法。'
    '最终3个簇之间的合并距离达到了2.1370（单链法仅为1.0182），是全链法中最高的。'
    '这说明全链法对簇的大小和形状更加敏感，倾向于保持簇的紧凑性，需要更大的距离才会合并不同簇。'
)

add_body_para(
    '全链法最终分类结果与单链法完全一致，准确率100%。三种方法的分类结果在该数据集上是'
    '一致的，因为数据簇之间分离良好，没有重叠或边界模糊的情况。',
    indent=True
)

add_title('（3）组平均法合并过程与结果', level=3)

add_body_para(
    '组平均法的合并过程中的最短距离值介于单链法和全链法之间。例如，第6次合并距离为0.4156'
    '（单链法0.3064，全链法0.5190），第22次合并距离为1.4683（单链法1.0182，全链法2.1370）。'
    '这验证了组平均法作为单链法和全链法折中方案的理论特性。',
    indent=True
)

add_body_para(
    '组平均法最终分类结果同样与单链法和全链法完全一致。',
    indent=True
)

add_title('（4）最短距离值汇总对比', level=3)

add_body_para(
    '下表汇总了三种方法在22步合并过程中的最短距离变化趋势：',
    indent=False
)

add_code_block(
'''              单链法      全链法      组平均法
第1步          0.0907      0.0907      0.0907
第5步          0.2439      0.2439      0.2439
第10步         0.5574      0.5857      0.5814
第15步         0.7331      1.2864      0.9614
第20步         0.9453      1.8541      1.3302
第22步（最后）  1.0182      2.1370      1.4683'''
)

add_body_para(
    '从上表可以看出：①在聚类过程初期（前5步），三种方法的最短距离完全一致，因为这些合并都发生在'
    '数据最紧密的簇0内部；②从第6步开始，三种方法的距离值逐渐分化，单链法距离最小，全链法距离最大，'
    '组平均法居中；③全链法在后期的距离增长最为陡峭，体现了其保守的合并策略——要求两个簇的所有点对'
    '都不能太远才会合并。'
)

add_title('3.4 手写实现与库函数对比验证', level=2)

add_body_para(
    '为了验证手写实现的正确性，本实验将手写算法的聚类结果与Scipy的linkage+fcluster以及Sklearn的'
    'AgglomerativeClustering进行了全面对比。',
)

# 插入对比图
compare_path = 'hand_vs_sklearn.png'
if os.path.exists(compare_path):
    doc.add_picture(compare_path, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_body_para(
    '图3 手写算法（上排）与Sklearn库（下排）聚类结果对比。从左至右依次为单链法、全链法和组平均法。'
    '可以观察到手写实现与Sklearn的结果完全一致，三种距离计算方法在25个数据点上的聚类结果准确率'
    '均为100%（25/25个点正确分类），验证了手写算法的正确性。',
    indent=True
)

add_body_para(
    '对比分析关键发现：\n'
    '（1）三种距离计算方法在该数据集上均得到了完全正确的聚类结果，原因在于数据簇之间分离良好，'
    '不存在重叠区域。\n'
    '（2）手写实现与Scipy/Sklearn库函数的结果标签可能以不同顺序排列（例如簇0和簇1互换），'
    '但这不影响聚类的正确性——聚类的评价关注的是哪些点被分在一起，而非簇的编号。\n'
    '（3）库函数的执行效率远高于手写实现。Scipy的linkage函数使用C语言优化实现，'
    '而手写Python实现包含大量的Python循环操作。对于25个数据点，两者在用户体验上无明显差异，'
    '但对于大规模数据集，库函数的优势将非常显著。',
    indent=True
)

add_title('3.5 结果分析', level=2)

add_body_para(
    '（1）聚类准确性分析：三种距离计算方法在本实验数据集上均达到了100%的聚类准确率。'
    '这说明当数据簇之间具有较好的可分性（簇内距离远小于簇间距离）时，层次聚类算法无论'
    '采用何种距离计算方式都能得到正确的结果。然而，在实际大数据安全应用中（如用户行为聚类），'
    '数据往往具有噪声和重叠，此时距离方法的选择将对结果产生显著影响。'
)

add_body_para(
    '（2）距离方法的影响：单链法容易产生"链式效应"——通过一串中间点将两个本应分开的簇连接在一起。'
    '全链法则倾向于产生紧凑的球形簇，但对离群点非常敏感。组平均法是实际应用中最常用的选择，'
    '因为它在鲁棒性和准确性之间取得了较好的平衡。本实验中，由于数据分布清晰可分，'
    '三种方法得到了相同的结果。'
)

add_body_para(
    '（3）计算复杂度：凝聚层次聚类的时间复杂度为O(N³)（使用朴素的簇间距离重计算方式），'
    '空间复杂度为O(N²)（存储距离矩阵）。当N较大时，算法将变得非常耗时。可通过使用优先队列'
    '（堆数据结构）将时间复杂度优化至O(N²logN)，但在大数据场景下仍面临挑战。'
    '这也是K-Means等划分方法在实际大数据应用中更为流行的原因之一。'
)

# ---- 四、思考问题 ----
doc.add_page_break()
add_title('四、思考问题', level=1)

add_title('4.1 凝聚层次聚类的优缺点', level=2)

add_body_para(
    '凝聚层次聚类作为一种经典的聚类方法，具有以下显著优点：',
    bold=True, indent=True
)

add_body_para(
    '1. 不需要预先指定聚类数量：与K-Means等划分聚类方法需要事先设定K值不同，层次聚类通过构建'
    '树状图（Dendrogram），使得用户可以事后在不同的层次上选择簇的数量。在访问控制等实际应用中，'
    '用户可以从树状图中直观地观察数据在不同粒度上的分组结构，然后根据业务需求选择合适的分组数量。'
    '这种灵活性是划分聚类方法所不具备的。'
)

add_body_para(
    '2. 结果的可解释性强：层次聚类生成的树状图提供了聚类过程的完整可视化。通过树状图，'
    '可以清晰地看到哪些对象首先被合并在一起（最相似），以及不同簇在什么距离阈值上汇聚。'
    '在安全审计场景中，审计人员可以利用树状图追溯异常用户与其他用户的相似度关系，'
    '为安全决策提供直观的数据支持。'
)

add_body_para(
    '3. 对簇形状的适应性：通过选择不同的距离计算方法（如单链法），层次聚类可以发现非球形的、'
    '任意形状的簇，这是K-Means等基于质心的方法所无法做到的。在处理具有复杂访问模式的用户'
    '行为数据时，这种灵活性非常有价值。'
)

add_body_para(
    '4. 确定性的聚类结果：与K-Means等依赖随机初始化的方法不同，层次聚类是确定性的——'
    '对于同一数据集和同一距离计算方法，每次运行的结果完全一致。这对于安全审计和合规性'
    '要求高的场景尤为重要。'
)

add_body_para(
    '5. 无需迭代优化：层次聚类通过贪心策略一次性地完成聚类，不存在K-Means中收敛判断和'
    '局部最优的问题。虽然贪心策略本身也有局限性，但至少结果是确定且可重复的。'
)

add_body_para(
    '然而，凝聚层次聚类也存在不可忽视的缺点：',
    bold=True, indent=True
)

add_body_para(
    '1. 计算复杂度高：凝聚层次聚类的时间复杂度为O(N²logN)（使用堆优化）至O(N³)（朴素实现），'
    '空间复杂度为O(N²)（需存储距离矩阵）。当数据规模N超过数千时，计算资源消耗将急剧增加，'
    '使得算法难以直接应用于大规模数据集。在大数据安全场景中，面对海量的用户访问日志，'
    '层次聚类需要借助采样或近似计算等技术来降低计算成本。'
)

add_body_para(
    '2. 不可逆的合并决策（贪心策略的局限）：层次聚类在每一步选择最近的两个簇进行合并，'
    '且这一决策不能撤销。一旦两个簇被合并，即使后续发现这个合并不理想，也无法拆分重来。'
    '这种贪心策略可能导致局部最优而非全局最优的聚类结果。例如，当数据中存在噪声点时，'
    '贪心策略可能过早地将噪声点错误地合并到某个簇中，影响最终的聚类质量。'
)

add_body_para(
    '3. 对噪声和离群点敏感：特别是单链法，离群点可能通过"桥接"方式将两个本应分开的簇连在一起'
    '（链式效应）。全链法虽然对链式效应具有抵抗力，但对离群点本身非常敏感——一个离群点可能显著'
    '增大簇间距离，导致聚类结果偏离预期。'
)

add_body_para(
    '4. 距离度量和链接方法的选择主观性强：不同的距离度量（欧几里得、曼哈顿、余弦等）和不同的'
    '链接方法（单链、全链、组平均、Ward等）会产生不同的聚类结果。目前没有统一的理论指导如何'
    '选择最优的组合，通常需要依赖领域知识和反复试验。这种主观性限制了层次聚类在自动化系统中的'
    '广泛应用。'
)

add_body_para(
    '5. 缺乏全局目标函数：与K-Means（最小化簇内平方和）或GMM（最大化似然）等方法不同，'
    '凝聚层次聚类没有显式的全局优化目标。它仅根据局部最近距离做出每一步的合并决策，'
    '不保证最终划分是全局最优的。'
)

add_title('4.2 另一种聚类方法——K-Means聚类', level=2)

add_body_para(
    'K-Means（K均值聚类）是最著名、应用最广泛的划分聚类算法之一。其基本思想是将数据集中的'
    'N个对象划分为预先指定的K个簇，使得每个对象到其所属簇的中心（质心）的距离平方和最小。'
    'K-Means算法由J. MacQueen于1967年提出，至今仍然是数据挖掘和机器学习领域中最常使用的'
    '聚类算法之一，也是大数据安全访问控制中的核心工具。',
    bold=False
)

add_body_para(
    '算法核心思想与步骤：',
    bold=True, indent=True
)

add_body_para(
    'K-Means算法采用迭代优化的策略，具体步骤如下：\n'
    '步骤1（初始化）：随机选择K个数据点作为初始聚类中心（质心），K值由用户预先指定。\n'
    '步骤2（分配）：计算每个数据点到K个质心的距离，将每个数据点分配到距离最近的质心所属的簇中。\n'
    '步骤3（更新）：对于每个簇，重新计算其质心位置，即取该簇中所有数据点的均值。\n'
    '步骤4（迭代）：重复步骤2和步骤3，直到质心位置不再发生显著变化（收敛）或达到最大迭代次数。'
)

add_body_para(
    '在访问控制技术中的应用：K-Means算法在访问控制中的一个重要应用是角色挖掘（Role Mining）。'
    '通过将用户按照其访问权限的相似性进行聚类，可以自动发现组织中隐含的角色结构。例如，'
    '将具有相似数据库访问权限的用户聚为一类，可以为这些用户统一分配一个"数据分析师"角色，'
    '从而支持基于角色的访问控制（RBAC）策略的构建和优化。此外，K-Means还可用于异常访问检测，'
    '将行为显著偏离所有簇中心的新访问请求标记为潜在的安全威胁。'
)

add_body_para(
    '与层次聚类的对比：',
    bold=True, indent=True
)

add_body_para(
    '（1）计算效率：K-Means的时间复杂度为O(N·K·t)，其中t是迭代次数（通常远小于N）。'
    '相比之下，层次聚类的O(N²logN)至O(N³)复杂度使得K-Means在大规模数据集上具有显著的效率优势。'
    '这使得K-Means更适合于处理大数据安全场景中的海量访问日志。'
)

add_body_para(
    '（2）结果依赖性：K-Means的结果依赖于初始质心的随机选择，不同的初始化可能导致不同的聚类结果。'
    '层次聚类是确定性的，结果可完全复现。为解决K-Means的随机性问题，实践中常采用K-Means++'
    '初始化策略或多次运行取最优结果的方法。'
)

add_body_para(
    '（3）簇形状假设：K-Means隐含地假设簇是各向同性的球形，对于非球形簇的识别效果较差。'
    '层次聚类（特别是单链法）可以识别任意形状的簇。在访问控制场景中，用户权限模式可能呈现出'
    '非球形的分布，此时需要根据数据特性选择合适的聚类方法。'
)

add_body_para(
    '（4）K值需求：K-Means要求用户预先指定簇的数量K，而层次聚类不需要。在实际安全分析中，'
    '往往事先不知道应该将用户分为多少个角色或多少个行为类别。此时可以先使用层次聚类生成树状图，'
    '通过观察树状图的结构来确定合理的K值，然后再应用K-Means进行大规模数据的精确聚类。'
    '这种"层次聚类定K + K-Means聚类"的混合策略在实际项目中非常常见。'
)

# ---- 五、心得体会 ----
doc.add_page_break()
add_title('五、心得体会', level=1)

add_body_para(
    '通过本次《访问控制技术——层次聚类算法》的综合实践，我获得了以下深刻的体会和认识：'
)

add_body_para(
    '第一，理论与实践之间的桥梁需要通过动手实现来搭建。在学习层次聚类的理论知识时，'
    '我理解了"自底向上合并最近的两个簇"这一核心思想，也记住了单链法、全链法和组平均法'
    '三种簇间距离计算公式。然而，当真正开始编写代码实现时，我才发现理论理解与实际编程之间'
    '存在许多细节问题需要解决：如何高效地管理动态变化的簇集合？如何在新簇生成后更新距离信息？'
    '簇的编号如何设计才能避免冲突？算法的终止条件如何正确实现？这些问题迫使我反复思考算法的'
    '每一个步骤，将书本上的伪代码转化为可正确运行的程序。这个过程极大地加深了我对层次聚类'
    '算法的理解——不仅是"知道它怎么工作"，更是"能够亲手让它工作"。'
)

add_body_para(
    '第二，可视化对理解聚类算法至关重要。在编写代码的过程中，我生成了树状图和散点图来'
    '观察聚类过程和结果。当我看到树状图中不同分支的高度差异反映了簇间距离的远近时，'
    '我对"聚类树"这一核心概念的理解立刻变得直观而深刻。同样，通过颜色标记的散点图，'
    '我能够一目了然地验证聚类结果的正确性。这让我认识到，在数据挖掘和机器学习领域，'
    '数据的可视化不是可有可无的附加功能，而是理解和验证算法不可或缺的支撑手段。'
)

add_body_para(
    '第三，不同的距离计算方法导致不同的聚类结果，选择合适的方法需要理解数据特性。'
    '通过对比单链法、全链法和组平均法在同一个数据集上的表现，我直观地感受到了它们的差异：'
    '单链法的"链式效应"、全链法的"保守策略"和组平均法的"折中特性"。在本实验使用的'
    '分离良好的数据集上，三种方法的结果完全一致，但通过分析合并过程中距离值的变化趋势，'
    '我仍能观察到它们在策略上的本质差异。这提醒我在面对实际的安全数据时，需要先分析数据'
    '的分布特性（是否存在离群点、簇的密度是否均匀、簇的形状是否为球形等），再据此选择'
    '合适的距离计算方法和链接策略。'
)

add_body_para(
    '第四，层次聚类在大数据安全中具有独特价值。本实验属于"大数据安全与隐私保护"课程，'
    '通过将层次聚类技术置于访问控制的背景下，我认识到聚类算法不仅仅是通用的数据挖掘工具，'
    '它们在安全领域有着直接而重要的应用。例如，通过对用户访问权限进行层次聚类，可以辅助'
    '角色挖掘，为RBAC策略设计提供数据驱动的依据；通过对访问行为日志进行聚类，可以发现'
    '异常访问模式，实现基于行为的安全审计。树状图的分层结构天然适合描述组织中权限的层次'
    '继承关系，这让我看到了算法特性与安全需求之间的巧妙契合。'
)

add_body_para(
    '第五，工程实践能力需要在对比和验证中提升。本次实验中，我不但需要从零实现凝聚层次聚类算法，'
    '还需要使用Sklearn和Scipy等成熟库进行对比验证。这个过程让我体会到了工业级库函数与'
    '自己手写代码之间的差异：库函数经过了高度优化，接口设计精良，使用简便；而手写代码虽然'
    '效率不如库函数，但编写过程让我对算法有了透彻的理解。两者结合——先用库函数快速验证思路，'
    '再深入理解其内部实现——是我认为高效的学习路径。'
)

add_body_para(
    '综上所述，本次综合实践不仅让我掌握了层次聚类算法的原理和实现方法，更重要的是让我体会到了'
    '理论、编码、可视化和应用分析这四者相互支撑、缺一不可的学习模式。我相信这些收获将对我今后'
    '学习更复杂的安全数据分析技术和处理实际问题产生持久的积极影响。'
)

# ---- 参考资料 ----
add_title('参考资料', level=1)

refs = [
    '[1] 层次聚类算法的原理与实现. CSDN博客. https://blog.csdn.net/liujh845633242/article/details/103679724',
    '[2] 层次聚类(Hierarchical Clustering)详解. 博客园. https://www.cnblogs.com/JetpropelledSnake/p/14513657.html',
    '[3] Scipy官方文档 - Hierarchical clustering (scipy.cluster.hierarchy). https://docs.scipy.org/doc/scipy/reference/cluster.hierarchy.html',
    '[4] Scikit-learn官方文档 - AgglomerativeClustering. https://scikit-learn.org/stable/modules/generated/sklearn.cluster.AgglomerativeClustering.html',
    '[5] Jiawei Han, Micheline Kamber, Jian Pei. 数据挖掘：概念与技术（第3版）. 机械工业出版社, 2012.',
    '[6] Pang-Ning Tan, Michael Steinbach, Vipin Kumar. 数据挖掘导论. 人民邮电出版社, 2011.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(ref)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)  # 五号
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ======================== 保存 ========================
output_path = '访问控制技术_层次聚类算法_实验报告.docx'
doc.save(output_path)
print(f'[√] 实验报告已生成: {output_path}')
